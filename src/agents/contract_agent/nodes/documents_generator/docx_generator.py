import os
import base64
import re
from io import BytesIO

from libs.logger import LoggerFactory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .documents_templates import CONTRACT_TEMPLATES

from ....utils import _normalize_space


logger = LoggerFactory.get_logger(
    name="ContractAgentDocumentGeneratorDocxNode",
    logs_path=os.getenv("LOGS_DIR"),
    log_file=os.getenv("LOGS_FILE") if os.getenv("MODE") != "DEBUG" else None,
)


def _add_inline_markdown(paragraph, text: str):
    """
    Очень простой inline-парсер для **bold** и *italic*.
    Достаточно для нормального DOCX без зависимости от pandoc.
    """
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _markdown_to_docx_bytes(markdown: str, document_title: str) -> bytes:
    doc = Document()

    # Базовая типографика
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document_title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    for line in markdown.splitlines():
        stripped = line.strip()

        if not stripped:
            continue

        # H1-H6
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=min(level, 4))
            continue

        # Маркеры списков
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_markdown(p, bullet_match.group(1).strip())
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_markdown(p, numbered_match.group(1).strip())
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        _add_inline_markdown(p, stripped)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def contract_docx_generation_node(state):
    logger.info("Start...")
    contract_type = state.get("contract_type")
    markdown = _normalize_space(state.get("generated_markdown", ""))

    document_title = CONTRACT_TEMPLATES[contract_type]["document_title"]
    docx_bytes = _markdown_to_docx_bytes(markdown, document_title)

    state["generated_docx_base64"] = base64.b64encode(docx_bytes).decode("utf-8")
    logger.debug(f"Got result docx document [:100]: {state["generated_docx_base64"][:max(100, len(state["generated_docx_base64"]))]}")
    logger.info("Finish")
    return state