"""
Генерация DOCX-документа из текста искового заявления.
"""
import os
import re
import base64
import io
import unicodedata
from datetime import datetime
from typing import Any

from logger import LoggerFactory

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = LoggerFactory.get_logger(
    name=__name__,
    logs_path=os.getenv("LOGS_DIR"),
    log_file=os.getenv("LOGS_FILE") if os.getenv("MODE") != "DEBUG" else None,
)


def generate_docx_bytes(document_text: str, metadata: dict[str, Any] | None = None) -> str:
    """
    Генерирует DOCX-файл из текста и возвращает base64-строку.

    Args:
        document_text: Текст искового заявления
        metadata: Метаданные (истец, ответчик и т.д.)

    Returns:
        Base64-строка DOCX-файла
    """
    logger.info("Generating DOCX from text (length: %d chars)", len(document_text))

    metadata = metadata or {}

    # Создаём документ
    doc = Document()

    # Настройка полей (2 см со всех сторон)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)    # ~2 см
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)   # ~3 см (для подшивки)
        section.right_margin = Inches(0.59)  # ~1.5 см

    # Парсим текст и форматируем
    _add_formatted_content(doc, document_text)

    # Добавляем метаданные как свойства документа (с ограничением 255 символов)
    if metadata:
        core_props = doc.core_properties
        core_props.title = "Исковое заявление"

        # Извлекаем короткие имена
        plaintiff_short = _extract_short_name(metadata.get('plaintiff', 'Не указан'), max_length=80)
        defendant_short = _extract_short_name(metadata.get('defendant', 'Не указан'), max_length=80)

        # Формируем subject (максимум 255 символов)
        subject = f"Истец: {plaintiff_short} / Ответчик: {defendant_short}"
        if len(subject) > 255:
            subject = subject[:252] + "..."

        core_props.subject = subject
        core_props.keywords = "исковое заявление, судебный документ"

    # Сохраняем в BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Конвертируем в base64
    docx_bytes = buffer.read()
    docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

    logger.info("DOCX generated, size: %d bytes, base64 length: %d", len(docx_bytes), len(docx_base64))

    return docx_bytes


def _extract_short_name(full_info: str, max_length: int = 80) -> str:
    """
    Извлекает короткое имя из полной информации.

    Примеры:
        "Иванов Иван Иванович, паспорт..." → "Иванов Иван Иванович"
        "ООО «Ромашка», ИНН 1234..." → "ООО «Ромашка»"
    """
    if not full_info:
        return "Не указан"

    # Берём первую строку до первой запятой или переноса
    lines = full_info.split('\n')
    first_line = lines[0].strip()

    parts = first_line.split(',')
    name = parts[0].strip() if parts else first_line

    # Обрезаем до max_length
    if len(name) > max_length:
        name = name[:max_length - 3] + "..."

    return name


_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_INLINE_MD_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _parse_line(raw_line: str) -> tuple[str, bool]:
    """Снимает префикс markdown-заголовка (# …). Возвращает (текст, был_заголовок)."""
    stripped = raw_line.strip()
    match = _MD_HEADING_RE.match(stripped)
    if match:
        return match.group(2).strip(), True
    return stripped, False


def _is_document_title(text: str) -> bool:
    upper = text.upper()
    return (
        "ИСКОВОЕ ЗАЯВЛЕНИЕ" in upper
        or "ДОСУДЕБНАЯ ПРЕТЕНЗИЯ" in upper
        or upper.startswith("ПРЕТЕНЗИ")
    )


def _is_section_heading(text: str, was_md_heading: bool) -> bool:
    if was_md_heading:
        return True
    if len(text) < 4 or len(text) > 120:
        return False
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    upper_count = sum(1 for c in letters if c.isupper() or c == "Ё")
    return upper_count / len(letters) >= 0.85


def _is_numbered_heading(text: str) -> bool:
    return bool(text) and text[0].isdigit() and "." in text[:5]


def _add_inline_runs(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    size: Pt = Pt(14),
) -> None:
    """Разбирает **жирный** и *курсив* в runs python-docx."""
    for part in _INLINE_MD_RE.split(text):
        if not part:
            continue
        run_bold = bold
        run_italic = False
        content = part
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            content = part[2:-2]
            run_bold = True
        elif (
            part.startswith("*")
            and part.endswith("*")
            and len(part) > 2
            and not part.startswith("**")
        ):
            content = part[1:-1]
            run_italic = True
        run = paragraph.add_run(content)
        run.font.name = "Times New Roman"
        run.font.size = size
        if run_bold:
            run.bold = True
        if run_italic:
            run.italic = True


def _add_formatted_content(doc: Document, text: str) -> None:
    """
    Добавляет текст в документ с форматированием.

    Правила:
    - Реквизиты (строки до заголовка) → справа, Times New Roman 12pt
    - Заголовок «ИСКОВОЕ ЗАЯВЛЕНИЕ» → по центру, жирный, 14pt
    - Основной текст → Times New Roman 14pt, междустрочный интервал 1.5
    """
    lines = text.split("\n")
    in_header = True

    for line in lines:
        line_stripped, was_md_heading = _parse_line(line)

        if not line_stripped:
            doc.add_paragraph()
            continue

        if _is_document_title(line_stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_runs(p, line_stripped, bold=True, size=Pt(14))
            in_header = False
            continue

        if in_header:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_inline_runs(p, line_stripped, size=Pt(12))
            continue

        if _is_section_heading(line_stripped, was_md_heading):
            p = doc.add_paragraph()
            _add_inline_runs(p, line_stripped, bold=True, size=Pt(14))
            continue

        if _is_numbered_heading(line_stripped):
            p = doc.add_paragraph()
            _add_inline_runs(p, line_stripped, bold=True, size=Pt(14))
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, line_stripped, size=Pt(14))
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Inches(0.5)


def save_docx_file(docx_bytes: bytes, path: str) -> None:
    with open(path, "wb") as f:
        f.write(docx_bytes)


def _party_slug(party_info: str | None, fallback: str) -> str:
    if not party_info or not str(party_info).strip():
        return fallback
    text = str(party_info).strip().split("\n")[0].split(",")[0].strip()
    text = re.sub(r"\s+", "_", text)
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'[<>:"/\\|?*]', "", text)
    text = text.strip("._ ")
    if len(text) > 48:
        text = text[:48].rstrip("._ ")
    return text or fallback


def _sanitize_filename_stem(stem: str, max_len: int = 180) -> str:
    stem = unicodedata.normalize("NFKC", stem)
    stem = re.sub(r'[<>:"/\\|?*]', "", stem)
    stem = re.sub(r"_+", "_", stem).strip("._ ")
    if not stem:
        return "document"
    if len(stem) > max_len:
        stem = stem[:max_len].rstrip("._ ")
    return stem


def build_docx_filename(
    document_type: str,
    plaintiff_info: str | None = None,
    defendant_info: str | None = None,
    generated_at: datetime | None = None,
) -> str:
    """
    Имя файла: {иск|претензия}_{истец}_{ответчик}_{YYYY-MM-DD_HH-MM-SS}.docx
    """
    kind = "претензия" if document_type == "complaint" else "иск"
    plaintiff = _party_slug(plaintiff_info, "истец")
    defendant = _party_slug(defendant_info, "ответчик")
    ts = (generated_at or datetime.now()).strftime("%Y-%m-%d_%H-%M-%S")
    stem = _sanitize_filename_stem(f"{kind}_{plaintiff}_{defendant}_{ts}")
    return f"{stem}.docx"


def resolve_unique_docx_path(directory: str, filename: str) -> str:
    path = os.path.join(directory, filename)
    if not os.path.exists(path):
        return path
    stem, ext = os.path.splitext(filename)
    n = 2
    while True:
        candidate = os.path.join(directory, f"{stem}_{n}{ext}")
        if not os.path.exists(candidate):
            return candidate
        n += 1